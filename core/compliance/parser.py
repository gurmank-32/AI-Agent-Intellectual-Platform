from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Union

from pydantic import BaseModel

# Matches patterns like:
#   "1. Title"  /  "1) Title"  /  "Section 1"  /  "CLAUSE 1"  /  "ARTICLE I"
#   or a line that is entirely UPPERCASE (heading)
_CLAUSE_SPLIT_RE = re.compile(
    r"""
    (?:^|\n)                           # start of text or newline
    (?:
        (?P<numbered>\d+)[.)]\s+       # "1. " or "1) "
      | Section\s+(?P<section>\d+)     # "Section 3"
      | CLAUSE\s+(?P<clause>\d+)       # "CLAUSE 5"
      | ARTICLE\s+(?P<article>[IVXLCDM\d]+)  # "ARTICLE IV" or "ARTICLE 4"
      | (?P<heading>[A-Z][A-Z\s]{3,})  # ALL-CAPS heading (≥4 chars)
    )
    """,
    re.VERBOSE | re.MULTILINE,
)

_ROMAN_MAP = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}


class Clause(BaseModel):
    number: int
    title: str
    content: str


class ParsedDocument(BaseModel):
    text: str
    clauses: list[Clause]


def _roman_to_int(s: str) -> int:
    total = 0
    prev = 0
    for ch in reversed(s.strip().upper()):
        val = _ROMAN_MAP.get(ch, 0)
        if val < prev:
            total -= val
        else:
            total += val
        prev = val
    return total


def _safe_int(value: str) -> int:
    try:
        return int(value)
    except ValueError:
        return _roman_to_int(value)


def parse_pdf(source: Union[str, Path, bytes, io.BytesIO]) -> ParsedDocument:
    from PyPDF2 import PdfReader

    if isinstance(source, (str, Path)):
        reader = PdfReader(str(source))
    elif isinstance(source, bytes):
        reader = PdfReader(io.BytesIO(source))
    else:
        reader = PdfReader(source)

    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)

    full_text = "\n".join(pages)
    clauses = extract_clauses(full_text)
    return ParsedDocument(text=full_text, clauses=clauses)


def parse_docx(source: Union[str, Path, bytes, io.BytesIO]) -> ParsedDocument:
    from docx import Document

    if isinstance(source, (str, Path)):
        doc = Document(str(source))
    elif isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    clauses = extract_clauses(full_text)
    return ParsedDocument(text=full_text, clauses=clauses)


def parse_document(
    source: Union[str, Path, bytes, io.BytesIO],
    filename: str = "",
) -> ParsedDocument:
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return parse_pdf(source)
    if name_lower.endswith(".docx"):
        return parse_docx(source)

    if isinstance(source, (str, Path)):
        name_lower = str(source).lower()
        if name_lower.endswith(".pdf"):
            return parse_pdf(source)
        if name_lower.endswith(".docx"):
            return parse_docx(source)

    raise ValueError(
        f"Unsupported file type: '{filename or source}'. Only PDF and DOCX are supported."
    )


def extract_clauses(text: str) -> list[Clause]:
    """Extract clauses from lease document text.

    Ported from the old `document_parser.py` logic:
    1) Try splitting by numbered clauses first (e.g. `1.`, `2)`).
    2) If that yields <= 1 clause, fall back to splitting by paragraphs.

    Additionally, when falling back, we also recognize ALL-CAPS headings
    and `CLAUSE X` markers as split points.
    """
    clauses: list[Clause] = []

    if not text or not text.strip():
        return []

    # 1) Try numbered clauses first (ported behavior)
    parts = re.split(r"\n\s*\d+[\.\)]\s+", text)
    for i, part in enumerate(parts):
        if not part.strip():
            continue
        lines = part.strip().split("\n")
        title = (lines[0][:100] if lines else f"Clause {i + 1}").strip()
        content = "\n".join(lines).strip()
        if content:
            clauses.append(Clause(number=i + 1, title=title, content=content))

    # 2) If no meaningful numbered splits, use paragraph fallback.
    if len(clauses) <= 1:
        # Split on double-newlines, then (optionally) split each paragraph further.
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        if not paragraphs:
            return [Clause(number=1, title="Full Document", content=text.strip())]

        # Recognize ALL-CAPS headings and "CLAUSE X" markers as internal split points.
        marker_split_re = re.compile(
            r"\n\s*(?:"
            r"\d+[\.\)]\s+|"  # numbered (within paragraph)
            r"CLAUSE\s+\d+\s+|"  # CLAUSE X
            r"[A-Z][A-Z\s]{3,}\s*\n|"  # ALL CAPS heading
            r"Section\s+\d+\s+|"  # Section N
            r"ARTICLE\s+\d+\s+"  # Article N
            r")"
        )

        expanded: list[Clause] = []
        for i, para in enumerate(paragraphs):
            subparts = [p.strip() for p in marker_split_re.split(para) if p.strip()]
            if not subparts:
                continue
            for j, sp in enumerate(subparts):
                lines = sp.split("\n")
                title = (lines[0][:100] if lines else para[:100]).strip()
                expanded.append(
                    Clause(number=len(expanded) + 1, title=title, content=sp)
                )

        clauses = expanded or [
            Clause(
                number=1,
                title=(paragraphs[0][:100]).strip(),
                content="\n\n".join(paragraphs),
            )
        ]

    return clauses
